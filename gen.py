from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

D = dict(
    fio="Абдельнур Алекс Фатех",
    uni="АНОО ВО «Университет «Сириус»",
    spec="10.02.04 Обеспечение информационной безопасности телекоммуникационных систем",
    course="4", group="К0409-22",
    sup_uni="Торшин Д. В.", sup_uni_f="Торшин Дмитрий Владимирович",
    sup_org="Оглоблин А. В.", sup_org_f="Оглоблин Александр Викторович",
    org="ООО «ЦОС»", addr="пгт. Сириус, Триумфальный проезд, д. 1",
    ds="09.02.2026", de="12.04.2026",
    theme="Разработка программного модуля для защиты веб-приложений от основных типов угроз"
)

def S(doc):
    for s in doc.sections:
        s.top_margin=Cm(2);s.bottom_margin=Cm(2);s.left_margin=Cm(3);s.right_margin=Cm(1.5)
    st=doc.styles['Normal']
    st.font.name='Times New Roman';st.font.size=Pt(14)
    st.paragraph_format.line_spacing=1.5
    st.paragraph_format.first_line_indent=Cm(1.25)
    st.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    st.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

def cp(doc,t,sz=14,b=False):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent=Pt(0)
    r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(sz);r.bold=b
    return p

def jp(doc,t):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing=1.5;p.paragraph_format.first_line_indent=Cm(1.25)
    r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(14)
    return p

def hp(doc,t):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(12);p.paragraph_format.space_after=Pt(6)
    p.paragraph_format.first_line_indent=Pt(0)
    r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(14);r.bold=True

def gen_report():
    doc=Document();S(doc)
    for _ in range(3):doc.add_paragraph()
    cp(doc,D['uni']);doc.add_paragraph()
    cp(doc,"ОТЧЁТ",18,True);cp(doc,"о прохождении преддипломной практики",16)
    doc.add_paragraph()
    cp(doc,"по направлению подготовки");cp(doc,D['spec'],14,True)
    doc.add_paragraph();cp(doc,"Тема: "+D['theme'])
    doc.add_paragraph()
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.RIGHT;p.paragraph_format.first_line_indent=Pt(0)
    r=p.add_run("Выполнил: студент "+D['course']+" курса\nгруппы "+D['group']+"\n"+D['fio']+"\n\n")
    r.font.name='Times New Roman';r.font.size=Pt(14)
    r=p.add_run("Руководитель от университета:\n"+D['sup_uni']+"\n\n")
    r.font.name='Times New Roman';r.font.size=Pt(14)
    r=p.add_run("Руководитель от организации:\n"+D['sup_org'])
    r.font.name='Times New Roman';r.font.size=Pt(14)
    doc.add_paragraph();doc.add_paragraph();cp(doc,"Сириус, 2026")
    doc.add_page_break()
    hp(doc,"РЕФЕРАТ")
    jp(doc,"Отчёт по преддипломной практике, период: "+D['ds']+" — "+D['de']+", организация: "+D['org']+".")
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.LEFT;p.paragraph_format.first_line_indent=Pt(0)
    r=p.add_run("Ключевые слова: ");r.font.name='Times New Roman';r.font.size=Pt(14);r.bold=True
    r=p.add_run("ВЕБ-БЕЗОПАСНОСТЬ, IDS/IPS, ЗАЩИТА ОТ АТАК, МОДУЛЬНАЯ АРХИТЕКТУРА");r.bold=False;r.font.name='Times New Roman';r.font.size=Pt(14)
    jp(doc,"Объект исследования — веб-приложения, подверженные SQL-инъекциям, XSS, RCE, LFI и другим атакам.")
    jp(doc,"Цель — разработка модуля sec для комплексной защиты веб-приложений на платформе AEngine.")
    jp(doc,"В ходе практики разработана архитектура модуля безопасности, реализованы подсистемы IDS/IPS, подписи кода, защиты ОС, сетевого анализа и мониторинга.")
    doc.add_page_break()
    hp(doc,"СОДЕРЖАНИЕ")
    toc=[("ВВЕДЕНИЕ","3"),("1. ХАРАКТЕРИСТИКА ОРГАНИЗАЦИИ","4"),("2. АНАЛИЗ СОСТОЯНИЯ ВОПРОСА","5"),("3. ПРОЕКТИРОВАНИЕ МОДУЛЯ","7"),("4. РЕАЛИЗАЦИЯ И ТЕСТИРОВАНИЕ","10"),("ЗАКЛЮЧЕНИЕ","14"),("СПИСОК ИСТОЧНИКОВ","15")]
    for t,pg in toc:
        p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.LEFT;p.paragraph_format.first_line_indent=Pt(0)
        r=p.add_run(t+" "+"."*40+" "+pg);r.font.name='Times New Roman';r.font.size=Pt(14)
    doc.add_page_break()
    hp(doc,"ВВЕДЕНИЕ")
    jp(doc,"Веб-приложения подвергаются растущему числу кибератак. Основные угрозы: SQLi, XSS, RCE, LFI, DoS.")
    jp(doc,"Актуальность обусловлена необходимостью комплексной защиты в реальном времени.")
    jp(doc,"Цель практики — разработка модуля sec для защиты веб-приложений от основных типов угроз.")
    jp(doc,"Задачи: изучить подходы к защите; проанализировать IDS/IPS; спроектировать архитектуру; реализовать модуль; протестировать.")
    doc.add_page_break()
    hp(doc,"1. ХАРАКТЕРИСТИКА ОРГАНИЗАЦИИ ПРАКТИКИ")
    jp(doc,"Практика проходила в "+D['org']+" ("+D['addr']+"). Организация специализируется на обеспечении информационной безопасности телекоммуникационных систем и разработке защитных решений.")
    jp(doc,"В организации функционирует центр информационной безопасности, занимающийся мониторингом, анализом угроз и реагированием на инциденты.")
    doc.add_page_break()
    hp(doc,"2. АНАЛИЗ СОСТОЯНИЯ ВОПРОСА")
    jp(doc,"Проведён анализ существующих решений: Snort, Suricata, ModSecurity, OWASP CRS. Выявлены преимущества и недостатки каждого подхода.")
    jp(doc,"Сигнатурные методы эффективны против известных атак, но требуют постоянного обновления баз. Поведенческий анализ обнаруживает zero-day угрозы, но даёт ложные срабатывания.")
    jp(doc,"Модульная архитектура позволяет гибко комбинировать методы защиты и масштабировать систему.")
    doc.add_page_break()
    hp(doc,"3. ПРОЕКТИРОВАНИЕ МОДУЛЯ БЕЗОПАСНОСТИ")
    jp(doc,"Модуль sec разработан для платформы AEngine с модульной архитектурой. Основные компоненты: подсистема обнаружения вторжений (intrusions.py), защита ОС (os_protect.py, sys_protect.py), сетевой анализ (net_analyzer.py), подпись кода (code_signer.py), кластеризация (cluster.py, auto_cluster.py), панель мониторинга (dashboard.py).")
    jp(doc,"Архитектура основана на принципе разделения ответственности: каждый подмодуль решает конкретную задачу безопасности.")
    doc.add_page_break()
    hp(doc,"4. РЕАЛИЗАЦИЯ И ТЕСТИРОВАНИЕ")
    jp(doc,"Реализована подсистема обнаружения вторжений с поддержкой сигнатурного анализа SQLi, XSS, RCE, LFI. База сигнатур хранится в signatures_db.json.")
    jp(doc,"Система подписи кода обеспечивает целостность файлов через хеширование и криптографическую верификацию.")
    jp(doc,"Модуль защиты ОС контролирует использование ресурсов, привилегии и конфигурации системы.")
    jp(doc,"Тестирование подтвердило эффективность обнаружения основных типов атак с минимальным уровнем ложных срабатываний.")
    doc.add_page_break()
    hp(doc,"ЗАКЛЮЧЕНИЕ")
    jp(doc,"В ходе преддипломной практики разработан программный модуль sec для защиты веб-приложений. Модуль включает IDS/IPS, подпись кода, защиту ОС, сетевой анализ и мониторинг.")
    jp(doc,"Результат — готовый к внедрению модуль, прошедший тестирование и демонстрирующий эффективность против основных типов кибератак.")
    doc.add_page_break()
    hp(doc,"СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    sources=["1. ГОСТ Р 57580.1-2017. Безопасность финансовых операций. Защита информации.","2. OWASP Top 10 Web Application Security Risks, 2021.","3. Snort Users Manual. Sourcefire, 2023.","4. ModSecurity Handbook. Ivan Ristic, 2013.","5. Документация платформы AEngine, 2026."]
    for s in sources:
        p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.LEFT;p.paragraph_format.first_line_indent=Cm(1.25)
        r=p.add_run(s);r.font.name='Times New Roman';r.font.size=Pt(14)
    doc.save("Отчет_практика.docx")
    print("Отчет создан: Отчет_практика.docx")

def gen_diary():
    doc=Document();S(doc)
    cp(doc,D['uni'],12);doc.add_paragraph()
    cp(doc,"ДНЕВНИК",16,True)
    cp(doc,"прохождения преддипломной практики",14)
    doc.add_paragraph()
    cp(doc,"обучающегося очной формы обучения "+D['course']+" курса группы "+D['group'],12)
    cp(doc,D['spec'],12)
    doc.add_paragraph()
    cp(doc,D['fio'],14,True)
    doc.add_paragraph()
    jp(doc,"Место практики: "+D['org']+", "+D['addr'])
    jp(doc,"Руководитель от организации: "+D['sup_org_f']+", "+D['sup_org'])
    jp(doc,"Сроки: с "+D['ds']+" по "+D['de'])
    doc.add_page_break()
    hp(doc,"График прохождения практики")
    tbl=doc.add_table(rows=1,cols=4);tbl.style='Table Grid'
    hdr=tbl.rows[0].cells;hdr[0].text="№";hdr[1].text="Наименование работ";hdr[2].text="Сроки";hdr[3].text="Примечание"
    for h in hdr:
        for p in h.paragraphs:
            for r in p.runs:r.bold=True;r.font.size=Pt(12)
    tasks=[
        ("1","Изучение архитектуры платформы AEngine и модуля sec","09.02-15.02",""),
        ("2","Анализ существующих IDS/IPS решений","16.02-22.02",""),
        ("3","Проектирование модульной архитектуры","23.02-01.03",""),
        ("4","Реализация подсистемы обнаружения вторжений","02.03-15.03",""),
        ("5","Разработка системы подписи кода","16.03-22.03",""),
        ("6","Реализация защиты ОС и сетевого анализа","23.03-29.03",""),
        ("7","Разработка панели мониторинга","30.03-05.04",""),
        ("8","Тестирование и отладка модуля","06.04-12.04",""),
    ]
    for n,t,s,p in tasks:
        row=tbl.add_row().cells;row[0].text=n;row[1].text=t;row[2].text=s;row[3].text=p
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:r.font.size=Pt(12)
    doc.add_page_break()
    hp(doc,"Ежедневные записи")
    tbl2=doc.add_table(rows=1,cols=3);tbl2.style='Table Grid'
    hdr2=tbl2.rows[0].cells;hdr2[0].text="Дата";hdr2[1].text="Выполненная работа";hdr2[2].text="Отметка"
    for h in hdr2:
        for p in h.paragraphs:
            for r in p.runs:r.bold=True;r.font.size=Pt(11)
    days=[
        ("09.02","Изучение документации AEngine","Выполнено"),
        ("10.02","Анализ структуры модуля sec","Выполнено"),
        ("11.02","Изучение module.json и архитектуры","Выполнено"),
        ("12.02","Обзор подсистемы intrusions.py","Выполнено"),
        ("13.02","Анализ signatures_db.json","Выполнено"),
        ("16.02","Сравнение Snort и Suricata","Выполнено"),
        ("17.02","Изучение ModSecurity и OWASP CRS","Выполнено"),
        ("18.02","Анализ поведенческих методов","Выполнено"),
        ("19.02","Обзор коммерческих WAF решений","Выполнено"),
        ("20.02","Сводный анализ IDS/IPS","Выполнено"),
        ("23.02","Проектирование архитектуры модуля","Выполнено"),
        ("24.02","Определение интерфейсов подмодулей","Выполнено"),
        ("25.02","Проектирование auth.py","Выполнено"),
        ("26.02","Проектирование системы подписи","Выполнено"),
        ("27.02","Схема кластеризации","Выполнено"),
        ("02.03","Реализация intrusions.py","Выполнено"),
        ("03.03","Добавление сигнатур SQLi","Выполнено"),
        ("04.03","Добавление сигнатур XSS","Выполнено"),
        ("05.03","Добавление сигнатур RCE, LFI","Выполнено"),
        ("06.03","Тестирование IDS","Выполнено"),
        ("09.03","Реализация code_signer.py","Выполнено"),
        ("10.03","Хеширование файлов","Выполнено"),
        ("11.03","Верификация подписи","Выполнено"),
        ("12.03","Тестирование подписи","Выполнено"),
        ("16.03","Реализация os_protect.py","Выполнено"),
        ("17.03","Реализация sys_protect.py","Выполнено"),
        ("18.03","Реализация net_analyzer.py","Выполнено"),
        ("19.03","Тестирование защиты ОС","Выполнено"),
        ("20.03","Тестирование сетевого анализа","Выполнено"),
        ("23.03","Реализация dashboard.py","Выполнено"),
        ("24.03","HTML шаблоны панели","Выполнено"),
        ("25.03","Авторизация в панели","Выполнено"),
        ("26.03","Интеграция компонентов","Выполнено"),
        ("30.03","Комплексное тестирование","Выполнено"),
        ("31.03","Исправление ошибок","Выполнено"),
        ("01.04","Оптимизация производительности","Выполнено"),
        ("02.04","Финальное тестирование","Выполнено"),
        ("06.04","Подготовка документации","Выполнено"),
        ("07.04","Оформление отчета","Выполнено"),
        ("08.04","Оформление дневника","Выполнено"),
        ("09.04","Подготовка к защите","Выполнено"),
        ("10.04","Финальная проверка модуля","Выполнено"),
        ("11.04","Сдача материалов","Выполнено"),
        ("12.04","Завершение практики","Выполнено"),
    ]
    for d,w,m in days:
        row=tbl2.add_row().cells;row[0].text=d;row[1].text=w;row[2].text=m
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:r.font.size=Pt(11)
    doc.add_page_break()
    hp(doc,"Отзыв руководителя")
    jp(doc,"За время прохождения практики студент "+D['fio']+" ознакомился с архитектурой платформы AEngine и модуля безопасности sec. Ответственно выполнил все поставленные задачи: изучил существующие IDS/IPS решения, спроектировал и реализовал модуль sec с подсистемами обнаружения вторжений, подписи кода, защиты ОС и мониторинга. Модуль прошёл тестирование и готов к внедрению.")
    doc.add_paragraph()
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.RIGHT;p.paragraph_format.first_line_indent=Pt(0)
    r=p.add_run("Руководитель от организации: _______________/"+D['sup_org']+"\n");r.font.name='Times New Roman';r.font.size=Pt(14)
    r=p.add_run("Руководитель от университета: _______________/"+D['sup_uni']+"\n");r.font.name='Times New Roman';r.font.size=Pt(14)
    r=p.add_run("Оценка: _______________");r.font.name='Times New Roman';r.font.size=Pt(14)
    doc.save("Дневник_практики.docx")
    print("Дневник создан: Дневник_практики.docx")

def gen_task():
    doc=Document();S(doc)
    cp(doc,D['uni'],12);doc.add_paragraph()
    cp(doc,"ИНДИВИДУАЛЬНОЕ ЗАДАНИЕ",16,True)
    cp(doc,"на преддипломную практику",14)
    doc.add_paragraph()
    cp(doc,D['fio'],14,True)
    cp(doc,D['spec'],12)
    doc.add_paragraph()
    jp(doc,"1. Тема практики: "+D['theme'])
    jp(doc,"2. Цель практики: Разработка программного модуля sec для защиты веб-приложений от основных типов угроз на платформе AEngine.")
    jp(doc,"3. Сроки: "+D['ds']+" — "+D['de'])
    doc.add_page_break()
    hp(doc,"План проведения практики")
    tbl=doc.add_table(rows=1,cols=3);tbl.style='Table Grid'
    hdr=tbl.rows[0].cells;hdr[0].text="№";hdr[1].text="Перечень заданий";hdr[2].text="Сроки"
    for h in hdr:
        for p in h.paragraphs:
            for r in p.runs:r.bold=True;r.font.size=Pt(12)
    tasks=[
        ("1","Изучение архитектуры платформы AEngine","09.02-15.02"),
        ("2","Анализ существующих IDS/IPS решений","16.02-22.02"),
        ("3","Проектирование модульной архитектуры sec","23.02-01.03"),
        ("4","Реализация подсистемы обнаружения вторжений","02.03-15.03"),
        ("5","Разработка системы подписи кода","16.03-22.03"),
        ("6","Реализация защиты ОС и сетевого анализа","23.03-29.03"),
        ("7","Разработка панели мониторинга безопасности","30.03-05.04"),
        ("8","Тестирование, отладка, оформление результатов","06.04-12.04"),
    ]
    for n,t,s in tasks:
        row=tbl.add_row().cells;row[0].text=n;row[1].text=t;row[2].text=s
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:r.font.size=Pt(12)
    doc.add_paragraph()
    jp(doc,"Дата выдачи: "+D['ds'])
    doc.add_paragraph()
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.LEFT;p.paragraph_format.first_line_indent=Pt(0)
    r=p.add_run("Руководитель от университета: _______________/"+D['sup_uni']+"\n");r.font.name='Times New Roman';r.font.size=Pt(14)
    r=p.add_run("Согласовано (организация): _______________/"+D['sup_org']+"\n");r.font.name='Times New Roman';r.font.size=Pt(14)
    r=p.add_run("Задание принял: _______________/"+D['fio']);r.font.name='Times New Roman';r.font.size=Pt(14)
    doc.save("Инд_задание.docx")
    print("Инд.задание создано: Инд_задание.docx")

if __name__=="__main__":
    gen_report()
    gen_diary()
    gen_task()
    print("Все документы созданы!")
